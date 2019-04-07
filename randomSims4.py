# -*- coding: utf-8 -*-
"""
Created on Tue Mar 26 12:53:04 2019

@author: Sarah
"""
import random
from docx import Document
from docx.shared import Pt

###############################################

def conventional_career():
    p.style.font.size = Pt(12)

    ran = random.randint(1,36)
    
    if ran == 1:
        p.add_run('Entertainer(Musician) - ').bold = True
        p.add_run('May write jingles\n').bold = False
    elif ran == 2:
        p.add_run('Entertainer(Comedian) - ').bold = True
        p.add_run('May perform comedy in space.\n').bold = False
    elif ran == 3:
        p.add_run('Writer(Author) - ').bold = True
        p.add_run('May publish all book genres.\n').bold = False
    elif ran == 4:
        p.add_run('Writer(Journalist) - ').bold = True
        p.add_run('May publish all book genres(self-publishing), and write articles.\n').bold = False
    elif ran == 5:
        p.add_run('Painter(Master of the Real) - ').bold = True
        p.add_run('May sell crafted paintings and call agent to collect money.\n').bold = False
    elif ran == 6:
        p.add_run('Painter(Patron of the Arts) - ').bold = True
        p.add_run('May sell crafted paintings and call agent to collect money.\n').bold = False
    elif ran == 7:
        p.add_run('Secret Agent(Diamond Agent) - ').bold = True
        p.add_run('No additional money-making opportunities.\n').bold = False
    elif ran == 8:
        p.add_run('Secret Agent(Villain) - ').bold = True
        p.add_run('No additional money-making opportunities.\n').bold = False
    elif ran == 9:
        p.add_run('Criminal(Boss) - ').bold = True
        p.add_run('May pickpocket sims.\n').bold = False
    elif ran == 10:
        p.add_run('Criminal(Oracle) - ').bold = True
        p.add_run('May pickpocket sims. May hack and create viruses.\n').bold = False
    elif ran == 11:
        p.add_run('Astronaut(Space Ranger) - ').bold = True
        p.add_run('May go on space missions with the rocket ship.\n').bold = False
    elif ran == 12:
        p.add_run('Astronaut(Smuggler) - ').bold = True
        p.add_run('May go on space missions with the rocket ship.\n').bold = False
    elif ran == 13:
        p.add_run('Culinary(Chef) - ').bold = True
        p.add_run('May publish cookbooks(self-publishing).\n').bold = False
    elif ran == 14:
        p.add_run('Culinary(Mixologist) - ').bold = True
        p.add_run('May publish bar guides(self-publishing).\n').bold = False
    elif ran == 15:
        p.add_run('Tech Guru(eSport Gamer) - ').bold = True
        p.add_run('No additional money-making opportunities.\n').bold = False
    elif ran == 16:
        p.add_run('Tech Guru(Start-up Entrepreneur) - ').bold = True
        p.add_run('No additional money-making opportunities.\n').bold = False
    elif ran == 17:
        p.add_run('Athlete(Bodybuilder) - ').bold = True
        p.add_run('No additional money-making opportunities.\n').bold = False
    elif ran == 18:
        p.add_run('Athlete(Pro Athlete) - ').bold = True
        p.add_run('No additional money-making opportunities.\n').bold = False
    elif ran == 19:
        p.add_run('Business(Management) - ').bold = True
        p.add_run('No additional money-making opportunities.\n').bold = False
    elif ran == 20:
        p.add_run('Business(Investor) - ').bold = True
        p.add_run('May invest in stocks.\n').bold = False
    elif ran == 21:
        p.add_run('Doctor - ').bold = True
        p.add_run('No additional money-making opportunities. This career counts as a profession.\n').bold = False
    elif ran == 22:
        p.add_run('Detective - ').bold = True
        p.add_run('No additional money-making opportunities. This career counts as a profession.\n').bold = False
    elif ran == 23:
        p.add_run('Scientist - ').bold = True
        p.add_run('No additional money-making opportunities. This career counts as a profession.\n').bold = False
    elif ran == 24:
        p.add_run('Critic(Art Critic) - ').bold = True
        p.add_run('No additional money-making opportunities.\n').bold = False
    elif ran == 25:
        p.add_run('Critic(Food Critic) - ').bold = True
        p.add_run('If you have rolled Homemade, your food critic(and only your food critic) may eat out for food review purposes. No additional money-making opportunities.\n').bold = False
    elif ran == 26:
        p.add_run('Politician(Charity Organizer) - ').bold = True
        p.add_run('No additional money-making opportunities.\n').bold = False
    elif ran == 27:
        p.add_run('Politician(Politician) - ').bold = True
        p.add_run('May ask for bribes.\n').bold = False
    elif ran == 28:
        p.add_run('Social Media(Internet Personality) - ').bold = True
        p.add_run('May advertise on your social media page.\n').bold = False
    elif ran == 29:
        p.add_run('Social Media(Public Relations) - ').bold = True
        p.add_run('No additional money-making opportunities.\n').bold = False
    elif ran == 30:
        p.add_run('Gardener(Botanist) - ').bold = True
        p.add_run('If you choose to also sell plants you are limited to one of each type.\n').bold = False
    elif ran == 31:
        p.add_run('Gardener(Flower Designer) - ').bold = True
        p.add_run('May only sell arrangements for income.\n').bold = False
    elif ran == 32:
        p.add_run('Stylist - ').bold = True
        p.add_run('I’m not sure if there are ways to make extra money with this career but if there are feel free to do so.\n').bold = False
    elif ran == 33:
        p.add_run('Trend Setter - ').bold = True
        p.add_run('I’m not sure if there are ways to make extra money with this career but if there are feel free to do so.\n').bold = False
    elif ran == 34:
        p.add_run('Actor - ').bold = True
        p.add_run('I know absolutely nothing about this career or the ways you can make money.  Normally with the “professions” such as scientist, you can only earn money from that job.  I’d say it’s the same but from what I’ve seen it’s not so easy to box in with the other hands-on careers.  For example, if you need to learn guitar and practice in public people automatically give tips.  I’ll update this when I see more of other’s playthroughs, for now, you have carte blanche to do whatever.  I’d love to hear your experience with this seeing-as-how I have no plans to get this pack anytime soon. \n').bold = False
    elif ran == 35:
        p.add_run('Military Cadet - ').bold = True
        p.add_run('This is an at-home career. \n').bold = False
    elif ran == 36:
        p.add_run('Military Covert Officer - ').bold = True
        p.add_run('This is also an at-home career.  You may bribe sims.\n').bold = False

###############################################
        
def unconventional_career():
    p.style.font.size = Pt(12)

    ran = random.randint(1,26)
    
    if ran == 1:
        p.add_run('Freelance Painter - ').bold = True
        p.add_run('Paint whatever you want whenever you want and sell it to collectors.\n').bold = False
    elif ran == 2:
        p.add_run('Freelance Author - ').bold = True
        p.add_run('Write and publish whatever you want whenever you want.\n').bold = False
    elif ran == 3:
        p.add_run('Freelance Musician - ').bold = True
        p.add_run('Play a musical instrument of your choice in public to earn money from tips. You may also sell jingles and license songs you’ve written. If you have City Living, singing counts as an instrument.\n').bold = False
    elif ran == 4:
        p.add_run('Gardener - ').bold = True
        p.add_run('Plant, harvest and sell produce and flowers of all kinds.\n').bold = False
    elif ran == 5:
        p.add_run('Fishersim - ').bold = True
        p.add_run('Sell whatever you can manage to catch.\n').bold = False
    elif ran == 6:
        p.add_run('Treasure Hunter - ').bold = True
        p.add_run('Dig up those little weird rocky-looking things and sell what you find. You may also sell the little treasures(aka, unwanted upgrade parts) that you dig out of your household plumbing, electronics, and appliances. Collecting and breeding frogs is also allowed, as is collecting and selling insects(OR). Selling collectibles you’ve traded for at the flea market is allowed.  The new dig spots (C&D) are allowed as well.\n').bold = False
    elif ran == 7:
        p.add_run('Carpenter - ').bold = True
        p.add_run('Use the wood crafting bench to make and sell furniture, as well as complete wood crafting projects.\n').bold = False
    elif ran == 8:
        p.add_run('Freelance Comedian - ').bold = True
        p.add_run('Use the microphone to perform comedy routines and publish comedy books(self-publishing).\n').bold = False
    elif ran == 9:
        p.add_run('Freelance Programmer - ').bold = True
        p.add_run('Use your programming skills on the computer for a variety of tasks, including creating apps, creating plugins and freelance work.\n').bold = False
    elif ran == 10:
        p.add_run('Space Explorer - ').bold = True
        p.add_run('Build your own rocket and take to the stars, selling what you bring back from your travels.\n').bold = False
    elif ran == 11:
        p.add_run('Professional Gamer - ').bold = True
        p.add_run('Create computer games(programming 9), Livestream and compete in gaming tournaments to earn money.\n').bold = False
    elif ran == 12:
        p.add_run('Cybercriminal - ').bold = True
        p.add_run('Create viruses and hack for money.\n').bold = False
    elif ran == 13:
        p.add_run('Professional Host/Hostess - ').bold = True
        p.add_run('Host social gatherings of all kinds and sell the rewards you earn for money.\n').bold = False
    elif ran == 14:
        p.add_run('Remedy Brewer - ').bold = True
        p.add_run('Brew herbal remedies and sell them for profit.\n').bold = False
    elif ran == 15:
        p.add_run('Freelance Photographer - ').bold = True
        p.add_run('Take and sell photographs.\n').bold = False
    elif ran == 16:
        p.add_run('Retail Owner - ').bold = True
        p.add_run('Manage a retail store that does not sell anything that’s covered by any of the other options above. For example, you could have a store that sells food you’ve cooked, but you could not sell raw produce at that store unless you had another sim in the household that rolled Gardener. Alternately, you may also choose to operate one of the vendor tables or display walls that come with City Living. This career counts as a profession unless you choose to operate a vendor table/display wall in lieu of a real store.\n').bold = False
    elif ran == 17:
        p.add_run('DJ - ').bold = True
        p.add_run('Work as a DJ to earn tips, and license your mixtapes.\n').bold = False
    elif ran == 18:
        p.add_run('Restaurateur - ').bold = True
        p.add_run('Manage your own restaurant. This career counts as a profession.\n').bold = False
    elif ran == 19:
        p.add_run('Animal Trainer - ').bold = True
        p.add_run('Adopt or Create a pet with the Prowler (c) or Hunter (d) trait and send them out to bring back loot for you.  It may be possible to also make money off of having a trained pet perform tricks.\n').bold = False
    elif ran == 20:
        p.add_run('Veterinarian  - ').bold = True
        p.add_run('May also craft pet treats to sell.\n').bold = False
    elif ran == 21:
        p.add_run('Archaeologist - ').bold = True
        p.add_run('Excavate and use the archaeology table to authenticate relics.  You may also write archaeology books and verify items mailed to you.  Note:  Only items you find through excavation can be sold.  If you choose to explore temples you can never sell the treasures.  You can, on the other hand, keep them to display in your home.\n').bold = False
    elif ran == 22:
        p.add_run('Treasure Hunter - ').bold = True
        p.add_run('Go on adventures to find simoleons, Omiscan treasures, relics, artifacts, and fossils to sell or keep.  Note:  Only items you find within temples or through pop-up quests can be sold.  If you choose to do archaeology such as excavating and using the table you can never sell the items.  You can, on the other hand, keep them to display in your home.\n').bold = False
    elif ran == 23:
        p.add_run('Flower Arranging - ').bold = True
        p.add_run('Craft flower arrangements to sell.  You may sell these via inventory, store, or yard sale table if this is possible.  You may have a small garden with flowers only but may not sell the flowers in any way other than through arrangements.\n').bold = False
    elif ran == 24:
        p.add_run('Bee Keeper - ').bold = True
        p.add_run('This was a requested unconventional career.  You can get as many hives as your heart desires.  Feel free to reroll this one.  I know a lot of you don’t have room for all those bee boxes or enjoy having uncomfortably stung sims all day.\n').bold = False
    elif ran == 25:
        p.add_run('Video Producer - ').bold = True
        p.add_run('Use the video station to record, edit, and add effects to your videos, then upload them.  You may also stream using the drone.\n').bold = False
    elif ran == 26:
        p.add_run('Music Producer - ').bold = True
        p.add_run('Use the Music Station to produce, customize, and release customized music tracks.\n').bold = False

###############################################

def midlife_crisis(num):
    if num == 1:
        p.add_run('* Re-roll Miscellaneous Fun(complete within 24 hours) – Your Miscellaneous Fun roll will change. It takes effect immediately, with weekly tasks starting the following Sunday.\n').bold = False
    elif num == 2:
        p.add_run('* Randomize a trait(complete within 24 hours) – Randomly determine which of your sim’s traits will change, then use a cheat to remove that trait, replacing it with another that has been randomly determined.\n').bold = False
    elif num == 3:
        p.add_run('* Re-roll Career(complete within 24 hours) – Re-roll your sim’s career. This means a new primary career if it’s the heir, and a new secondary career if it’s a spouse or spare.\n').bold = False
    elif num == 4:
        p.add_run('* Make a major purchase(complete within 24 hours) – Buy something shiny and expensive(relative to how much wealth you have) for your house.\n').bold = False
    elif num == 5:
        p.add_run('* Have a romantic crisis(start within 48 hours) – Your sim’s love life is thrown into turmoil. Perhaps they meet a new lover, cheat on their spouse, or simply have a falling out with a long-time partner. This must begin within 48 hours but doesn’t have to resolve ever.\n').bold = False
    elif num == 6:
        p.add_run('* Change style to recapture youth(complete within 24 hours) – Your sim must change their style(mirror and dresser, CAS cheats if necessary) to recapture their youth. Perhaps even get a tattoo. They can change back after a period of a few days if they want.\n').bold = False
    elif num == 7:
        p.add_run('* Gain a new skill(complete within 1 week) – Your sim must gain at least 5 skill levels in a skill they haven’t used much previously. This counts from the level they’re currently at, so if a sim has level 2 charisma, they must attain level 7.\n').bold = False
    elif num == 8:
        p.add_run('* Have a child(complete within 1 week) – Your sim must either have a biological child or adopt a child. Either way, it will be raised with the rest of the generation’s children. Add one child to the number of children you will have this generation.\n').bold = False
    elif num == 9:
        p.add_run('* Randomize Aspiration(complete within 24 hours) – Randomly pick a new aspiration for your sim, throwing out any that are incompatible with your other rolls. If you also have to re-roll your sim’s career, use the new career when determining aspiration compatibility.\n').bold = False
    elif num == 10:
        p.add_run('* Adopt a pet(complete within 24 hours) – Adopt a pet through an agency or befriend a stray and bring them home.\n').bold = False
    elif num == 11:
        p.add_run('* Plus One(complete within 1 week) A new helper is added to the household.  It can be any sim that has not previously lived in the household.  If you rolled single for this generation it can be a spouse or live-in partner.  Perhaps it is a long lost family member?  You may roll a secondary career for the sim but their traits should be randomized.  If you have a full house already reroll. \n').bold = False
    

###############################################
    
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'

###############################################

#Gender Law 1-3
p = doc.add_paragraph()
p.style.font.bold = True
p.add_run('Gender Law: ').font.size = Pt(16)
p.style.font.size = Pt(12)


ran = random.randint(1,3)
if ran == 1:
    p.add_run('Matriarchy: ').bold = True
    p.add_run('The Founder must be female. Only girls are eligible to be named heir unless there are no female children, at which point boys become eligible for that generation.\n').bold = False
    
elif ran == 2:
    p.add_run('Patriarchy: ').bold = True
    p.add_run('The Founder must be male. Only boys are eligible to be named heir unless there are no male children, at which point girls become eligible for that generation.\n').bold = False
elif ran == 3:
    p.add_run('Equality: ').bold = True
    p.add_run('The Founder may be of either gender. Both boys and girls are eligible for the title of heir.\n').bold = False
    
###############################################
    
#Bloodline Law 1-3
p = doc.add_paragraph()
p.style.font.bold = True
p.add_run('Bloodline Law: ').font.size = Pt(16)


ran = random.randint(1,3)
if ran == 1:
    p.add_run('Traditional: ').bold = True
    p.add_run('Children who are naturally born from the previous generation are eligible to be named heir. Adopted children are ineligible to be named heir unless there are no naturally born children, at which point they become eligible for that generation.\n').bold = False
elif ran == 2:
    p.add_run('Modern: ').bold = True
    p.add_run('Both Naturally born and adopted children are eligible to be named heir.\n').bold = False
elif ran == 3:
    p.add_run('Foster: ').bold = True
    p.add_run('Children who are adopted are eligible to be named heir. Naturally born children are not eligible to be named heir unless there are no adopted children, at which point they become eligible for that generation\n').bold = False
    
###############################################
    
#Heir Law 1-7
p = doc.add_paragraph()
p.style.font.bold = True
p.add_run('Heir Law: ').font.size = Pt(16)


ran = random.randint(1,7)
if ran == 1:
    p.add_run('First Born: ').bold = True
    p.add_run('The oldest, by order of joining the family, eligible living child is named heir.\n').bold = False
elif ran == 2:
    p.add_run('Last Born: ').bold = True
    p.add_run('The youngest, by order of joining the family, eligible living child is named heir.\n').bold = False
elif ran == 3:
    p.add_run('Living Will: ').bold = True
    p.add_run('The eligible child with the highest friendly relationship score with their previous-generation’s parent will be named heir.\n').bold = False
elif ran == 4:
    p.add_run('Merit: ').bold = True
    p.add_run('The child with the most fully completed aspirations will be named heir. If there is a tie, the child with the highest level in a single skill will become heir from among the children who are tied.\n').bold = False
elif ran == 5:
    p.add_run('Strength: ').bold = True
    p.add_run('The first born eligible child becomes heir by default… but the tile can be forcefully taken from them if an eligible sibling beats them in a fight. That sibling may have their title taken, (or taken back) if they lose a fight to another eligible sibling.\n').bold = False
elif ran == 6:
    p.add_run('Random: ').bold = True
    p.add_run('The title of heir is randomly selected from the pool of all eligible children. Every time the eligible pool changes size, The heir must be re-rolled using the new pool\n').bold = False
elif ran == 7:
    p.add_run('Exemplar: ').bold = True
    p.add_run('At the beginning of the challenge, name a single trait. This trait must be one of your founder’s three traits.. Any eligible heir that has this trait will gain the title of heir. If a single generation has no children with this trait or more than one exemplar, follow the First Born rule.\n').bold = False
    
###############################################

#Which Gen is it?
p = doc.add_paragraph()
runner = p.add_run('Generation #: \n')
runner.underline = True
runner.font.size = Pt(16)

###############################################

#Marital Status 1-25
p = doc.add_paragraph()
p.style.font.bold = True
p.add_run('Marital Status: ').font.size = Pt(16)

helpers = 0
ran = random.randint(1,25)
if (ran >= 1) & (ran <= 5):
    p.add_run('Single - ').bold = True
    p.add_run('Your heir may not have a live-in partner after the first child is born. Any extra sims must be moved out of the house by 24 hours after the first child is born. A male heir may move in a female partner for the course of any pregnancies, but the 24-hour rule applies and she may not contribute financially. Partners that have moved out are not allowed to lurk around the house to help out with toddler care and must be removed if they will not leave.\n').bold = False
elif (ran >= 6) & (ran <= 15):
    helpers = 1
    p.add_run('Couple - ').bold = True
    p.add_run('Your heir must obtain a romantic live-in partner. They do not have to be married.\n').bold = False
elif ran == 16:
    p.add_run('Mixed Single - ').bold = True
    p.add_run('Children must be a mix of naturally born and adopted. If you have rolled to have only one child, you may have a second to satisfy this parameter.\n').bold = False
elif ran == 17:
    helpers = 1
    p.add_run('Mixed Couple- ').bold = True
    p.add_run('Your heir must obtain a romantic live-in partner. They do not have to be married. Children must be a mix of naturally born and adopted. If you have rolled to have only one child, you may have a second to satisfy this parameter.\n').bold = False
elif ran == 18:
    helpers = 2
    p.add_run('Second Chance - ').bold = True
    p.add_run('Your heir must obtain two live-in partners during the course of the generation. At least one child must be born or adopted with the first partner.\n').bold = False
elif ran == 19:
    helpers = 1
    p.add_run('Single w/ Help - ').bold = True
    p.add_run('A non-romantic helper sim must live in the house. The helper sim can be anything you desire, from a friend of the family to a spare who decides to stick around.\n').bold = False
elif ran == 20:
    helpers = 1
    p.add_run('Mixed Single w/ Help - ').bold = True
    p.add_run('Children must be a mix of naturally born and adopted. If you have rolled to have only one child, you may have a second to satisfy this parameter. A non-romantic helper sim must live in the house. The helper sim can be anything you desire, from a friend of the family to a spare who decides to stick around.\n').bold = False
elif ran == 21:
    helpers = 2
    p.add_run('Couple w/ Help - ').bold = True
    p.add_run('Your heir must obtain a romantic live-in partner. They do not have to be married. A non-romantic helper sim must live in the house. The helper sim can be anything you desire, from a friend of the family to a spare who decides to stick around.\n').bold = False
elif ran == 22:
    helpers = 2
    p.add_run('Mixed Couple w/ Help - ').bold = True
    p.add_run('Children must be a mix of naturally born and adopted. If you have rolled to have only one child, you may have a second to satisfy this parameter. Your heir must obtain a romantic live-in partner. They do not have to be married. A non-romantic helper sim must live in the house. The helper sim can be anything you desire, from a friend of the family to a spare who decides to stick around.\n').bold = False
elif ran == 23:
    helpers = 2
    p.add_run('Single w/ 2 Help – ').bold = True
    p.add_run('The same as single with help, except there are two helpers. The helpers may be romantically involved with each other.\n').bold = False
elif ran == 24:
    helpers = 2
    p.add_run('Mixed Single w/ 2 Help - ').bold = True
    p.add_run('Children must be a mix of naturally born and adopted. If you have rolled to have only one child, you may have a second to satisfy this parameter. The same as single with help, except there are two helpers. The helpers may be romantically involved with each other.\n').bold = False
elif ran == 25:
    helpers = 3
    p.add_run('Full House - ').bold = True
    p.add_run('The same as single with 2 help, except there are three helpers. \n').bold = False
    
############################################### 
    
#Number of Children 1-10
p = doc.add_paragraph()
p.style.font.bold = True
p.add_run('Number of Children: ').font.size = Pt(16)

ran = random.randint(1,10)

if (ran >=1) & (ran <= 2):
    p.add_run('1\n').bold = False
elif (ran >=3) & (ran <= 6):
    p.add_run('2\n').bold = False
elif (ran >=7) & (ran <= 8):
    p.add_run('3\n').bold = False
elif ran == 9:
    p.add_run('5\n').bold = False
elif ran == 10:
    p.add_run('5\n').bold = False
    
###############################################
    
#Primary Career 1-10
p = doc.add_paragraph()
p.style.font.bold = True
p.add_run('Primary Career: ').font.size = Pt(16)

ran = random.randint(1,10)

if (ran >=1) & (ran <= 6):
    p.add_run('Conventional Career: \n').font.size = Pt(14)
    conventional_career()
elif (ran >=7) & (ran <= 9):
    p.add_run('Unconventional Career: \n').font.size = Pt(14)
    unconventional_career()
elif ran == 10:
    p.add_run('Career Hopper – ').bold = True
    p.add_run('Must roll for five conventional careers. You may jump between these careers as you wish, so long as you have spent a day at work in at least three by the time your sim becomes a full adult, and all five by the time your sim becomes an elder.\n').bold = False
    conventional_career()
    conventional_career()
    conventional_career()
    conventional_career()
    conventional_career()
    
###############################################

#Secondary Career 1-10
if helpers > 0:
    p = doc.add_paragraph()
    p.style.font.bold = True
    p.add_run('Secondary Career: \n').font.size = Pt(16)
    
    for i in range(helpers):
        ran = random.randint(1,10)
        
        if (ran >=1) & (ran <= 3):
            p.add_run('%d' %(i+1) + ' Conventional Career: \n').font.size = Pt(14)
            conventional_career()
        elif (ran >=4) & (ran <= 8):
            p.add_run('%d' %(i+1) + ' Unconventional Career: \n').font.size = Pt(14)
            unconventional_career()
        elif ran == 9:
            p.add_run('%d' %(i+1) + ' Career Hopper – ').bold = True
            p.add_run('Must roll for five conventional careers. You may jump between these careers as you wish, so long as you have spent a day at work in at least three by the time your sim becomes a full adult, and all five by the time your sim becomes an elder.\n').bold = False
            conventional_career()
            conventional_career()
            conventional_career()
            conventional_career()
            conventional_career()
        elif ran == 10:
            p.add_run('%d' %(i+1) + ' Unemployed').bold = True
            
###############################################

#Generation Goals 1-26    
p = doc.add_paragraph()
p.style.font.bold = True
p.add_run('Generation Goals: ').font.size = Pt(16)

ran = random.randint(1,26)
if ran == 1:
    p.add_run('Perfect Careers - ').bold = True
    p.add_run('The heir, spouse, and helpers must reach level 10 in their careers. Re-roll if no conventional careers are rolled this generation.  You may also re-roll if you are doing career hopper. It might be possible to top the final career if you want to try but I’ve never been able to. \n').bold = False
elif ran == 2:
    p.add_run('Fulfilled - ').bold = True
    p.add_run('The heir and spouse(if applicable) must complete their first aspiration. If the spouse is accidentally too old when they move in to have time to complete an aspiration, they are not required to do so, but try to be reasonable with this provision.\n').bold = False
elif ran == 3:
    p.add_run('Perfect Children - ').bold = True
    p.add_run('All children born this generation must complete their childhood aspiration and receive an B grade in elementary school and an A grade in high school. They should also have a BFF made during childhood, and at least one boyfriend or girlfriend lasting at least 24 sim hours during their teen years. As toddlers, they must reach level 3 in all skills, including potty training, and raise at least one skill to level 5.  Each child must gain at least 3 positive character traits when aging to YA.\n').bold = False
elif ran == 4:
    p.add_run('Dependable - ').bold = True
    p.add_run('Your heir, spouse and helpers must complete all daily tasks for school and work, beginning in childhood or upon move-in to the house. You may use vacation days given by the game to skip school or work if you’re not going to be able to handle the task, but you may not skip if you have no vacation, or without using your vacation, or attend without having your daily task completed(even if you have saved vacation). Re-roll if no non-profession conventional careers are rolled this generation.\n').bold = False
elif ran == 5:
    p.add_run('Expansionist - ').bold = True
    p.add_run('You must construct a significant expansion to the house this generation. You may also add or change a single lot trait, if you wish. If Random is your miscellaneous fun this generation, the trait changed and new trait must be randomly determined.\n').bold = False
elif ran == 6:
    p.add_run('Collector - ').bold = True
    p.add_run('Your heir, spouse, helpers and children must complete and display one of the collections. All items must be found or grown that generation, and for plants you may not use seeds from previous generations…find your own! If you already completed part of a collection in a previous generation, you can choose that collection, but you must collect the objects you already have again with a sim from this generation. Collections and display requirements are as follows: To avoid confusion I’ve removed the very long and confusing list of collection types and how they should be displayed.  A collection is any of those you find by clicking the diamond in the personal inventory window.  You know you’ve completed it because a plaque will arrive in the mail.  Display however you’d like but they should be displayed somewhere through the entirety of the challenge.\n').bold = False
elif ran == 7:
    p.add_run('Deadbeat Parents - ').bold = True
    p.add_run('The heir, spouse, and helpers may not help their children this generation, apart from basic needs like food and interaction(autonomous chatting over dinner). This includes helping with homework, mentoring or encouraging children in skills, and reading books to them. For toddlers, it is permitted to provide basic care such as food and diaper changes, but try to avoid any parent-child interactions that raise a skill bar, such as flash cards, reading books, potty training, and so on. Some autonomous talking is okay – the parents are deadbeat, not abusive – but don’t deliberately train the communication skill. Children may teach themselves.  Each child should age up to YA with at least one negative character trait.\n').bold = False
elif ran == 8:
    p.add_run('Change of Scenery - ').bold = True
    p.add_run('It’s moving time! At some point, after the heir comes of age, you must either move to a new lot or completely demolish the entire house and start over from scratch on your current lot. Your new lot may not share any lot traits with your old lot. If Random is your miscellaneous fun this generation, the new lot’s traits must be randomly determined. Re-roll if it’s generation 1.\n').bold = False
elif ran == 9:
    p.add_run('Party King/Party Queen - ').bold = True
    p.add_run('Host every kind of party(house, dinner, wedding, birthday) with a gold medal in each. If you have more than four party types(through expansions, DLC or mods), roll to randomly determine the four party types that you will be attempting for this goal, re-rolling any duplicates.\n').bold = False
elif ran == 10:
    p.add_run('Idle Careers - ').bold = True
    p.add_run('May not make an effort to advance in your heir’s, helper’s and spouse’s careers, including school. If the career goals happen to line up with aspirational goals, that’s a happy coincidence, because you may advance in aspiration freely. Autonomous actions are also acceptable. If you’ve rolled an unconventional career that has no tasks that sims will perform autonomously, you may only perform tasks for it if your sim has a whim to do it or if they do them autonomously.  If you rolled for this generation before the heir aged to YA they may not have the Responsible trait!  If you rolled an interactive career just never choose the “Travel to Work” option – send them alone, as the work performance should, in theory, stay stagnant that way.\n').bold = False
elif ran == 11:
    p.add_run('Friendship is Forever - ').bold = True
    p.add_run('Beginning when the heir is a child, they must make one new good friend every week. This friendship must be maintained through subsequent weeks, and may only be ended by the death of the friend.\n').bold = False
elif ran == 12:
    p.add_run('Midlife Crisis - ').bold = True
    p.add_run('This generation, one of the heir, spouses(including second chances) or helpers will experience a randomly-generated midlife crisis event. Between when this generation is rolled and when the first of the eligible sims reach the adult life stage(not YA), determine who will go through the crisis. If you have not picked someone before the first sim reaches adulthood, or is moved in as an adult, that sim will be the one. Once the selected sim reaches adulthood(and not sooner!), roll three times to determine three aspects of the midlife crisis from the chart below, re-rolling if you receive a duplicate. Each aspect has a deadline which it must either be completed or started during. The goal is complete when all aspects have been completed.\n').bold = False
    midran1 = random.randint(1,11)
    midran2 = random.randint(1,11)
    while midran2 == midran1:
        midran2 = random.randint(1,11)
    midran3 = random.randint(1,11)
    while midran3 == midran2 | midran3 == midran1:
        midran3 = random.randint(1,11)
    midlife_crisis(midran1)
    midlife_crisis(midran2)
    midlife_crisis(midran3)
elif ran == 13:
    p.add_run('Memorial - ').bold = True
    p.add_run('The heir, along with any spouses or helpers, must be memorialized this generation, by a sim with maxed skill in the relevant crafting skill(painting, photography or writing). Memorials may consist of: a painted portrait, a photograph, a biography, or a book of life. Each memorial can be crafted by a different sim if you want, and they don’t have to all be the same type, though they should all be displayed together.\n').bold = False
elif ran == 14:
    p.add_run('Themed Display- ').bold = True
    p.add_run('This generation, the heir, spouses and helpers should create a themed display somewhere in their house or yard. The display must contain at least 10 items, plant and craft level excellent or higher, from at least two different collections or craft disciplines, and must contain at least one craft and one collected item. For example, a display with 3 plants, 5 photographs and 2 space rocks would be okay, but a display with 4 photographs and 6 paintings would not be, because there is no collected item. The theme is left up to you, examples include: space(telescope prints and photographs of the alien homeworld), botany(plants, photographs of plants, and microscope prints of plants), flowers(plants and paintings of plants in the wild) and animals(frogs, insects and photographs of insects in the wild).\n').bold = False
elif ran == 15:
    p.add_run('Haunted House - ').bold = True
    p.add_run('From the time the heir comes of age to the time the heir dies, three sims must die an unnatural death on your home lot. Unnatural death is defined as any death except by old age, and the sims may be members of your household(such as unneeded spouses, spares that are about to move out, or even the heir themselves) or guests. Furthermore, the graves must stay on your home lot, and the ghosts be allowed to roam the house freely. The graves may be removed after this generation’s heir dies.\n').bold = False
elif ran == 16:
    p.add_run('Best Club Ever - ').bold = True
    p.add_run('Your heir must start their own club from scratch this generation. To complete this goal, the following criteria must be met concurrently: \n* The club must have 8 members. Sims may not join the club unless they are already friends with the club leader, so get to recruiting!\n* The club must have a hangout built specially for them, either on your home lot or on a community lot. The hangout should contain enough seating for the entire club and any necessary items to complete club activities. If you build on a community lot from the map view, take note of how much money you “spent” on the hangout and subtract it from your family’s funds. \n* The club should have all ranks of one of the emotional vibes purchased, depending on the club’s focus. Happy counts. \n* The club door perk must be unlocked(since you have a hangout) \n* In addition, three other club spirit perks(hats, jackets, wall decorations, etc) must be unlocked and used.\n').bold = False
elif ran == 17:
    p.add_run('Domestic Dilemma - ').bold = True
    p.add_run('When this generation comes of age(or at the start of the game) you must either add a new “negative” trait or replace an existing trait(determine which of the three randomly) with a “negative” trait of your choice(unless you’ve rolled Random). This generation, you must find a creative way in your story to overcome said negative trait. For example, Quake Zone might be overcome by a sim maxing their writing skill and writing 5 nonfiction bestseller books about earthquake safety, or perhaps by a sim maxing their handiness skill and equipping every appliance/plumbing object in the house with the unbreakable upgrade. Be creative! When the negative trait has been overcome, you may replace it with a trait of your choice if you wish(again, unless you’ve rolled Random).\nNegative traits include:\nCursed\nFilthy\nGremlins\nGrody\nHaunted\nMean Vibe\nOn a Dark Ley Line\nQuake Zone\nVampire Nexus\nCreepy Crawlies\n').bold = False
elif ran == 18:
    p.add_run('My Honor Student(s) - ').bold = True
    p.add_run('In this generation complete every school project in every possible level from poor-excellent and create a prominent display area so that any guests to the house can marvel at your child’s intellect… even if the parent did most of it themselves.\n').bold = False
elif ran == 19:
    p.add_run('Pet Shrine - ').bold = True
    p.add_run('Adopt, create, or bring home a stray pet.  Through the generation build a shrine to the pet.  Example: A wall of Simstigram photos, their urn, favorite toys, collected feathers, collected presents unopened.\n').bold = False
elif ran == 20:
    p.add_run('Bad Investment - ').bold = True
    p.add_run('At the end of the previous generation the former heir made a bad investment and lost most of the families’ money!  Add their household funds and lot value then multiply that by 90%.  This is the amount you will need to subtract.  Once there are no funds left the “repo man” comes.  You will need to sell any and all household items until the debt is repaid.  You may choose to sell the house and start over if you’d like.\n').bold = False
elif ran == 21:
    p.add_run('Bring Me Back To Life - ').bold = True
    p.add_run('At some point during this generation, one of the family members teen-elder will have a terrible accident… of your choice of course.  After Grim has done his duty have the remaining family members work towards having the poor soul rejoin the family and resurrect them.  Note: If you won’t have at least one living caretaker left in the house re-roll.\n').bold = False
elif ran == 22:
    p.add_run('Do-Gooders - ').bold = True
    p.add_run('This family’s mantra is doing onto others.  And they do!  They volunteer together at least once a week.  They feed strays in the park.  They invite homeless sims to dinner and let them stay the night.  They regularly donate to charity.  Basically, any do-gooder thing you can think of!\n').bold = False
elif ran == 23:
    p.add_run('Holiday Fanatics - ').bold = True
    p.add_run('Must celebrate at least one holiday each sim week with 5 goals and all completed if possible depending on the sim’s likes. \n').bold = False
elif ran == 24:
    p.add_run('Scouting Family- ').bold = True
    p.add_run('Your family is into scouting.  Really into it.  Even though the parents can’t be scout leaders they take their children’s involvement seriously.  Any badge that can be earned with a parent tagging along they should.  Example: Taking the kids to Granite Falls for the nature badge.  \n').bold = False
elif ran == 25:
    p.add_run('Rich and Famous- ').bold = True
    p.add_run('This generation your family is all about the Fame!  Everyone in the household must do what they can to gain fame.  Once a week they should go somewhere to be seen by their adoring fans and sign autographs, have their picture taken, etc.  If they receive gifts they should be set aside in a shrine type area or a dedicated room to display their fame.  Duplicate gifts may be deleted but not sold.  \n').bold = False
elif ran == 26:
    p.add_run('Doomsday Preppers - ').bold = True
    p.add_run('At least one member of the family must solve the StrangerVille mystery and it’s aspiration.  If you want the entire family can do it.  This aspiration can be in addition to your regular one.  So if you have a musican doing Musical Genius you can switch back and forth.  Optional extra fun: Temporarily move the family to StrangerVille preferably to a bunker style house.  You can do this either until the mystery is solved or for the entirety of the generation.  You may change the lot traits.  When starting the next generation you may choose any lot.  You aren’t required to return to the previous legacy lot.  You may also stay on your StrangerVille lot.  Keep in mind if you choose to stay you can’t move again until rolling the “change of scenery” generational goal.  At the start of that next generation you may again choose new lot traits.  I’ve added this in case you wanted to temporarily use some of the traits that make things dirty or infested with spiders and such while living in a bunker. \n').bold = False
    
###############################################

#Miscellaneous Fun 1-30    
p = doc.add_paragraph()
p.style.font.bold = True
p.add_run('Miscellaneous Fun: ').font.size = Pt(16)

ran = random.randint(1,30)

if ran == 1:
    p.add_run('Random - ').bold = True
    p.add_run('Determine traits and aspirations(including the childhood aspiration) randomly for all sims of this generation(the heir and their siblings). Adult aspirations should be chosen randomly from the list of aspirations, discarding any which are incompatible with other parameters, as stated earlier in the rules. Helpers and Spouses moved in should have their adult aspiration randomized in a similar manner if you choose to change it. Any lot traits determined this generation, including initial move-in traits, must be random.\n').bold = False
elif ran == 2:
    p.add_run('Partier - ').bold = True
    p.add_run('Throw at least one party(of any type and rating) every week.\n').bold = False
elif ran == 3:
    p.add_run('Mischief Managed - ').bold = True
    p.add_run('Your heir must play a successful prank once every week from the time they age into a teen. Pranks include: making a prank call(mischief 3+), clogging drains at other lots(mischief 5+), calling in to play hooky(mischief 6+), sabotaging a rocket ship(mischief 8+), using a stink potion(mental 10+), sabotaging objects and sims(Tormentor reward trait), and any prank options performed with a partner in crime(mischief 4+).\n').bold = False
elif ran == 4:
    p.add_run('Moody - ').bold = True
    p.add_run('Once a week, your heir and spouse(if applicable) must spend at least 3 consecutive hours in a certain mood, randomly chosen from the following list: inspired, happy, angry, embarrassed, playful, flirty, sad, bored, confident, focused, tense. Heir and spouse may have different required moods.  Cannot use emotional aura items beyond viewing them to aid in establishing or continuing the mood.  Must participate in activities corresponding to the rolled mood for the 3 hour period such as jogging while energized or fulfilling mood whims.  If they slip out of the mood get them back into it.  Carl has a guide that has links to more information on achieving each emotion that should be very helpful: Sims 4 Emotions\n').bold = False
elif ran == 5:
    p.add_run('Homemade - ').bold = True
    p.add_run('May not get quick meals from the fridge, order pizza, eat from food stalls or dine out, consume plasma packs, and may only cook meals with ingredients if all ingredients are present. The pre-made toddler meals you can serve from the high chair are considered quick meals.\n').bold = False
elif ran == 6:
    p.add_run('Runs in the Family - ').bold = True
    p.add_run('Pick a trait from your heir(not a bonus or reward trait). Every child born to your heir must take this trait at the earliest opportunity.\n').bold = False
elif ran == 7:
    p.add_run('Perfect Attendance - ').bold = True
    p.add_run('From childhood onwards, your heir and helpers must never be more than one hour late to work or school. Similarly, they may not make use of the “leave early” function. A five strike allowance is provided for accidental lateness(like if you queue up actions and don’t notice that they didn’t leave for work), but intentional lateness for playing strategy or the fifth lateness will result in a loss of the challenge. This restriction also applies to the heir’s children and any spouses or helpers moved in later. Perfect attendance also applies to the use of vacation days, either while on vacation or for family leave.\n').bold = False
elif ran == 8:
    p.add_run('Fixer Upper - ').bold = True
    p.add_run('May not call the handyman or purchase new objects if they break this generation. You must fix them. I don’t care if you have no handiness skill, if you want to shower today you’ll get to work, won’t you?\n').bold = False
elif ran == 9:
    p.add_run('Half-Siblings - ').bold = True
    p.add_run('Sims born this generation must each be with a different partner. Re-roll if you’re only having one kid.\n').bold = False
elif ran == 10:
    p.add_run('Fighter - ').bold = True
    p.add_run('The heir, spouse or a helper(pick one and stick with it) must get in a fight at least once a week.\n').bold = False
elif ran == 11:
    p.add_run('Joker - ').bold = True
    p.add_run('Free re-roll of one category, now or in a future generation.\n').bold = False
elif ran == 12:
    p.add_run('Health Nut - ').bold = True
    p.add_run('This generation, sims may only eat healthy food.  In addition, the heir, spouse, helpers, and their children must each spend at least 3 hours each week performing physical activity.  Any activity that increases movement, motor, fitness or wellness skills count.  Note:  I found a list of the calorie counts of sim’s foods.  Might add a bit of fun.  Otherwise, I’d suggest sticking to salads and such.  Basically whatever is healthy in real life.  Example: Grilled fish is healthy but fish in a sauce is not.  Here’s the list: Sims 4 Food Calorie Values\n').bold = False
elif ran == 13:
    p.add_run('Gourmet - ').bold = True
    p.add_run('Once every week, the entire family(heir, spouse, and children) must find the time to sit down all at once and enjoy a gourmet meal prepared especially for the occasion. It doesn’t matter whether the heir or the spouse prepares the meal, but it must be from the gourmet cooking skill, not the regular cooking skill.  Add: If you have Parenthood have a child set the table as well.\n').bold = False
elif ran == 14:
    p.add_run('Hands Off! - ').bold = True
    p.add_run('You may not use the “hand of god” to clean up books and so on. Sims must clean up on their own. Selling things from their inventory is okay.\n').bold = False
elif ran == 15:
    p.add_run('Hobbyist - ').bold = True
    p.add_run('Your heir and spouse(if applicable) must each pick a hobby activity(such as painting, playing chess, jogging, fishing, etc) that is unrelated to their rolled career and devote at least 3 hours every sim week to practicing that activity. They may not use the activity to make money.\n').bold = False
elif ran == 16:
    p.add_run('My Precious - ').bold = True
    p.add_run('When your sim finds something they like, it’s hard for them to let it go. Pick one of the options below to build a collection over the course of your heir’s lifetime. The collection must be displayed somewhere on the legacy lot so it may not be kept in an inventory until the generation is complete.\nA) I finds it, it’s mine – At least once a week, your sim must go out and find a unique object to add to their collection(pick a single category: frogs, MySims Trophies, metals, crystals, elements, fossils, aliens, space rocks or fish). You may not display duplicates. If you have found one of every item and the generation isn’t complete yet, pick a new collection to begin.\nB) So bright, so beautiful, my precious – Once a week, your sim must purchase something beautiful and shiny from the decorations tab in buy mode. This should be something that exists only to look awesome and should be displayed all together(or as trophy centerpieces in each room or something), not decor that “blends in” such as toilet paper rolls in the bathroom or a pan rack in the kitchen. If your sim is insane and decides to build a “bathroom shrine” in their living room, that’s totally legit, but it should be done to add to the story, not just because you couldn’t be bothered to find a better idea than buying misc decorations for your bathrooms to fulfill this fun.\n').bold = False
elif ran == 17:
    p.add_run('Vacationer - ').bold = True
    p.add_run('At least one sim from the household must go on vacation at least once every sim week. If you don’t have the OR game pack use the following alternate rule since that sim can’t go on an actual vacation. They must use at least one vacation day from work and must have a day outing of some sort, lasting for most of the day(you know the kind you’d take in real life, leaving after breakfast and arriving back home in the evening), on their day off. This doesn’t have to be the same sim every week.\n').bold = False
elif ran == 18:
    p.add_run('Clubber - ').bold = True
    p.add_run('Starting from childhood or the time they join the household, the heir, and every spouse/helper must belong to a club, attending a meeting at least once a week. They do not have to belong to the same club, though they can if you want. If you choose to use a household club, please keep in mind the rules about how to avoid making clubs cheaty(#19 under the Rules), specifically the ones prohibiting having a constant club gathering running at all times. They must actually get together in a group and do a thing in order to have a club gathering, you can’t just run a gathering while they’re all sleeping or doing random tasks around the house.\n').bold = False
elif ran == 19:
    p.add_run('On Call - ').bold = True
    p.add_run('Starting when the heir is a teen, they(or a spouse or helper) must respond to at least one “hang-out” phone call every week. This includes birthday parties and invitations to community lots, but does not include requests to come to hang out at your home lot – you must actually leave your home lot for this to count.\n').bold = False
elif ran == 20:
    p.add_run('Town Explorer - ').bold = True
    p.add_run('Starting when the heir is a teen, they(or a spouse or helper) must visit a unique community lot every week, for at least 3 hours. The community lots may not be repeated to count for this, though you may visit them again just for fun.\n').bold = False
elif ran == 21:
    p.add_run('Regular Diner - ').bold = True
    p.add_run('Once a week, the entire household must go out to eat at a restaurant. All three courses and something to drink should be ordered.\n').bold = False
elif ran == 22:
    p.add_run('Festival Frequenter - ').bold = True
    p.add_run('Every time there is a festival, at least one household member must attend the festival for at least one sim hour. This doesn’t have to be the same household member each time.\n').bold = False
elif ran == 23:
    p.add_run('Flea Market Fanatic - ').bold = True
    p.add_run('Some sims just can’t pass up a great bargain. Why would you pay full price for anything? This generation, you may not purchase certain categories of items from the catalog. These items may only be purchased at the flea market or from unplayable sims selling their own goods on the street. Paintings left on public easels count.\nThe categories are:\nComfort(except beds)\nLighting\nDecorations(except mirrors)\n').bold = False
elif ran == 24:
    p.add_run('Sibling Rivalry - ').bold = True
    p.add_run('Whenever the parents aren’t within earshot the siblings can only interact using mean or mischievous interactions until one of them ages to young adult and they patch things up.\n').bold = False
elif ran == 25:
    p.add_run('Rainbow Brite - ').bold = True
    p.add_run('This generation your heir’s taste drastically changed!  Using the design tool in build/buy mode change every item possible to the rolled color.  For extra torture/fun dress the entire family in the chosen color as well!\n').bold = False
    color = random.randint(1,10)
    if color == 1:
        p.add_run('Red!\n').bold = True
    elif color == 2:
        p.add_run('Pink!\n').bold = True
    elif color == 3:
        p.add_run('Yellow!\n').bold = True
    elif color == 4:
        p.add_run('Green!\n').bold = True
    elif color == 5:
        p.add_run('Blue!\n').bold = True
    elif color == 6:
        p.add_run('Purple!\n').bold = True
    elif color == 7:
        p.add_run('White!\n').bold = True
    elif color == 8:
        p.add_run('Gray!\n').bold = True
    elif color == 9:
        p.add_run('Black!\n').bold = True
    elif color == 10:
        p.add_run('Brown!\n').bold = True
elif ran == 26:
    p.add_run('Exchange Student - ').bold = True
    p.add_run('When one of the children is a teenager they are invited to participate in an exchange program.  Pick any other household in the save that has another teenager and swap them in manage worlds.  Befriend the exchange student.  You can have them stay for only a few days or as long as a week.  After they have left invite them to visit once a week and remain friends with at least one family member.\n').bold = False
elif ran == 27:
    p.add_run('Foster Pets- ').bold = True
    p.add_run('From the start of this generation to the time the next generation’s heir takes over you must have a pet in the house that you are actively training daily to encourage someone to adopt it.  Once the pet has trained away at least 3 misbehaviors like drinking or rolling in pee puddles and has learned some tricks (if it’s a dog) you will choose a family to “adopt” them via manage worlds.  If you won’t have room for a pet see the rule regarding pets.\n').bold = False
elif ran == 28:
    p.add_run('Cultured - ').bold = True
    p.add_run('All family members should max the Selvadoradian culture skill, visit at least twice during this generation to immerse themselves.  Optional: Complete the relics collection.\n').bold = False
elif ran == 29:
    p.add_run('Seasonal “Fun” - ').bold = True
    p.add_run('Each season you must decorate the sim’s house & wardrobe according to the season, as well as do something (like a party or group outing) to celebrate the season.\n').bold = False
elif ran == 30:
    p.add_run('Crazed Fans - ').bold = True
    p.add_run('Your family is all about the rich and famous.  At least once a week they should seek out and fawn over the celebrity elite but should not go out of their way to attempt to become famous themselves.  If they gain fame inadvertently that’s fine.\n').bold = False

    
doc.save("randomroll.docx")